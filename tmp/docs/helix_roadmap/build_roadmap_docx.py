from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[3]
SOURCE = ROOT / "output" / "doc" / "helix_6_week_roadmap.md"
TARGET = ROOT / "output" / "doc" / "helix_6_week_roadmap.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_padding(cell, width: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for edge in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(width))
        node.set(qn("w:type"), "dxa")


def append_inline_runs(paragraph, text: str, *, bold: bool = False) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part[1:-1] if part.startswith("`") and part.endswith("`") else part)
        run.bold = bold
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Menlo"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Menlo")
            run.font.size = Pt(9.5)


def add_rich_paragraph(document: Document, text: str, *, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    match = re.match(r"^\*\*(.+?)\*\*(.*)$", text)
    if match:
        append_inline_runs(paragraph, match.group(1), bold=True)
        append_inline_runs(paragraph, match.group(2))
    else:
        append_inline_runs(paragraph, text)
    paragraph.paragraph_format.space_after = Pt(6)


def add_table(document: Document, table_lines: list[str]) -> None:
    rows = []
    for raw in table_lines:
        line = raw.strip()
        if not line or set(line.replace("|", "").replace("-", "").replace(" ", "")) == set():
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        rows.append(cells)

    if len(rows) < 2:
        return

    header = rows[0]
    body = rows[1:]
    table = document.add_table(rows=len(body) + 1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    column_widths = [Inches(0.85), Inches(1.7), Inches(2.65), Inches(2.3)]

    for row_index, values in enumerate([header] + body):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            cell.width = column_widths[min(col_index, len(column_widths) - 1)]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_padding(cell)
            if row_index == 0:
                set_cell_shading(cell, "D9E7F5")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            append_inline_runs(paragraph, value, bold=row_index == 0)


def build_doc() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(12)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            title = document.add_paragraph(style="Title")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            append_inline_runs(title, stripped[2:])
            title.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        if stripped.startswith("## "):
            document.add_heading(stripped[3:], level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            document.add_heading(stripped[4:], level=2)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(document, table_lines)
            continue

        if stripped.startswith("- "):
            bullet = document.add_paragraph(style="List Bullet")
            append_inline_runs(bullet, stripped[2:])
            bullet.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            candidate = lines[i].strip()
            if not candidate or candidate.startswith(("# ", "## ", "### ", "- ", "|")):
                break
            paragraph_lines.append(candidate)
            i += 1
        add_rich_paragraph(document, " ".join(paragraph_lines))

    footer_section = document.sections[-1]
    footer = footer_section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        "Generated and refreshed from repository inspection on 2026-03-18"
    )
    footer_run.italic = True
    footer_run.font.size = Pt(8.5)

    TARGET.parent.mkdir(parents=True, exist_ok=True)
    document.save(TARGET)


if __name__ == "__main__":
    build_doc()
